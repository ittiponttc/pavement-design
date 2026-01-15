import streamlit as st
import pandas as pd
from docx import Document

# =====================================================
# ค่าคงที่
# =====================================================
CM_TO_INCH = 1 / 2.54
MPA_TO_PSI = 145.038

# =====================================================
# ฐานข้อมูลวัสดุชั้นทาง (MR : MPa)
# =====================================================
MATERIAL_DB = {
    "ผิวทางลาดยาง PMA": {"E_default": 3700, "E_min": 2500, "E_max": 5000},
    "ผิวทางลาดยาง AC": {"E_default": 2500, "E_min": 1800, "E_max": 3500},
    "พื้นทางซีเมนต์ CTB": {"E_default": 1200, "E_min": 800, "E_max": 2000},
    "หินคลุกผสมซีเมนต์ (UCS 24.5 ksc)": {"E_default": 850, "E_min": 600, "E_max": 1200},
    "หินคลุก CBR 80%": {"E_default": 350, "E_min": 250, "E_max": 500},
    "ดินซีเมนต์ (UCS 17.5 ksc)": {"E_default": 350, "E_min": 200, "E_max": 600},
    "วัสดุหมุนเวียน (Recycling)": {"E_default": 850, "E_min": 600, "E_max": 1200},
    "รองพื้นทางวัสดุมวลรวม (CBR 25%)": {"E_default": 150, "E_min": 100, "E_max": 250},
    "วัสดุคัดเลือก ก": {"E_default": 76, "E_min": 50, "E_max": 120},
    "ดินถมคันทาง / ดินเดิม": {"E_default": 100, "E_min": 50, "E_max": 150},
}

# =====================================================
# ฟังก์ชัน: อัปเดต MR เมื่อเปลี่ยนชนิดวัสดุ
# =====================================================
def update_MR(i):
    mat = st.session_state[f"mat_{i}"]
    st.session_state[f"E_{i}"] = MATERIAL_DB[mat]["E_default"]

# =====================================================
# ตั้งค่าหน้าเว็บ
# =====================================================
st.set_page_config(page_title="Equivalent Modulus (Odemark)", layout="centered")
st.title("การคำนวณโมดูลัสเทียบเท่าของโครงสร้างทาง")
st.subheader("วิธี Odemark (1974)")
st.markdown("กรอกความหนาเป็น **เซนติเมตร (cm)** และค่า **MR เป็น MPa**")

# =====================================================
# เลือกจำนวนชั้น
# =====================================================
n_layers = st.slider("จำนวนชั้นโครงสร้างทาง", 1, 5, 3)

layers = []

# =====================================================
# กรอกข้อมูลแต่ละชั้น
# =====================================================
st.markdown("### ข้อมูลชั้นทาง")

for i in range(n_layers):

    # ตั้งค่าเริ่มต้นครั้งแรก (กัน key error)
    if f"mat_{i}" not in st.session_state:
        first_mat = list(MATERIAL_DB.keys())[0]
        st.session_state[f"mat_{i}"] = first_mat
        st.session_state[f"E_{i}"] = MATERIAL_DB[first_mat]["E_default"]

    with st.expander(f"ชั้นที่ {i+1}", expanded=True):

        col1, col2, col3 = st.columns(3)

        with col1:
            material = st.selectbox(
                "ชนิดวัสดุ",
                list(MATERIAL_DB.keys()),
                key=f"mat_{i}",
                on_change=update_MR,
                args=(i,)
            )

        with col2:
            thickness_cm = st.number_input(
                "ความหนา (ซม.)",
                min_value=0.1,
                value=10.0,
                step=0.5,
                key=f"h_{i}"
            )

        with col3:
            MR = st.number_input(
                "Modulus Resilient, MR (MPa)",
                min_value=10.0,
                step=50.0,
                key=f"E_{i}"
            )

        Emin = MATERIAL_DB[material]["E_min"]
        Emax = MATERIAL_DB[material]["E_max"]

        if not (Emin <= MR <= Emax):
            st.warning(
                f"ค่า MR = {MR:.0f} MPa อยู่นอกช่วงแนะนำ "
                f"({Emin} – {Emax} MPa)"
            )

        layers.append({
            "ชั้น": f"ชั้นที่ {i+1}",
            "ชนิดวัสดุ": material,
            "ความหนา (ซม.)": thickness_cm,
            "ความหนา (นิ้ว)": thickness_cm * CM_TO_INCH,
            "MR (MPa)": MR
        })

# =====================================================
# ปุ่มคำนวณ
# =====================================================
if st.button("คำนวณโมดูลัสเทียบเท่า"):

    sum_h = sum(l["ความหนา (ซม.)"] for l in layers)
    sum_h_E13 = sum(
        l["ความหนา (ซม.)"] * (l["MR (MPa)"] ** (1/3)) for l in layers
    )

    Eeq_MPa = (sum_h_E13 / sum_h) ** 3
    Eeq_psi = Eeq_MPa * MPA_TO_PSI

    df = pd.DataFrame(layers)

    st.markdown("### สรุปข้อมูลที่กรอก")
    st.dataframe(df, use_container_width=True)

    st.success(
        f"โมดูลัสเทียบเท่า (E_equivalent) = "
        f"{Eeq_psi:,.0f} psi ({Eeq_MPa:.1f} MPa)"
    )

    # =================================================
    # แสดง / ซ่อน วิธีการคำนวณ
    # =================================================
    if st.checkbox("แสดงวิธีการคำนวณ"):
        st.latex(
            r"E_{eq} = \left( \frac{\sum h_i E_i^{1/3}}{\sum h_i} \right)^3"
        )

        for l in layers:
            st.write(
                f"- {l['ชั้น']} : "
                f"h = {l['ความหนา (ซม.)']:.2f} cm, "
                f"MR = {l['MR (MPa)']:.1f} MPa, "
                f"MR¹ᐟ³ = {(l['MR (MPa)']**(1/3)):.3f}"
            )

        st.write(f"Σh = {sum_h:.2f} cm")
        st.write(f"Σ(h·MR¹ᐟ³) = {sum_h_E13:.2f}")

    # =================================================
    # สร้างรายงาน Word
    # =================================================
    doc = Document()
    doc.add_heading("การคำนวณโมดูลัสเทียบเท่าของโครงสร้างทาง", level=1)
    doc.add_paragraph("วิธี Odemark (1974)\n")

    doc.add_heading("ข้อมูลชั้นทาง", level=2)
    for l in layers:
        doc.add_paragraph(
            f"{l['ชั้น']} : {l['ชนิดวัสดุ']} | "
            f"h = {l['ความหนา (ซม.)']:.2f} cm "
            f"({l['ความหนา (นิ้ว)']:.2f} in), "
            f"MR = {l['MR (MPa)']:.1f} MPa"
        )
doc.add_heading("วิธีการคำนวณ", level=2)
doc.add_paragraph(
    "E_eq = ( Σ(h·E^(1/3)) / Σh )^3"
)
doc.add_paragraph(f"Σh = {sum_h:.2f} cm")
doc.add_paragraph(f"Σ(h·E^(1/3)) = {sum_h_E13:.2f}")
    doc.add_heading("ผลการคำนวณ", level=2)
    doc.add_paragraph(
        f"E_equivalent = {Eeq_psi:,.0f} psi ({Eeq_MPa:.1f} MPa)"
    )

    file_name = "Equivalent_Modulus_Odemark.docx"
    doc.save(file_name)

    with open(file_name, "rb") as f:
        st.download_button(
            "ดาวน์โหลดรายงาน (Word)",
            f,
            file_name=file_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
