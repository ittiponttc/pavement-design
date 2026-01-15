"""
โปรแกรมออกแบบและตรวจสอบความหนาถนนคอนกรีต (Rigid Pavement)
ตามวิธี AASHTO 1993
รองรับทั้ง JPCP (Jointed Plain Concrete Pavement) และ CRCP (Continuously Reinforced Concrete Pavement)

พัฒนาสำหรับใช้ในการเรียนการสอน
ภาควิชาครุศาสตร์โยธา มหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ
"""

import streamlit as st
import math
from io import BytesIO
from datetime import datetime

# ============================================================
# ส่วนที่ 1: ค่าคงที่และตารางอ้างอิง AASHTO 1993
# ============================================================

# ตารางค่า ZR (Standard Normal Deviate) ตามระดับความเชื่อมั่น
ZR_TABLE = {
    50: -0.000,
    60: -0.253,
    70: -0.524,
    75: -0.674,
    80: -0.841,
    85: -1.037,
    90: -1.282,
    91: -1.340,
    92: -1.405,
    93: -1.476,
    94: -1.555,
    95: -1.645,
    96: -1.751,
    97: -1.881,
    98: -2.054,
    99: -2.327
}

# ค่า Load Transfer Coefficient (J) ตามประเภทถนนและการถ่ายแรง
# อ้างอิง: AASHTO 1993 Guide, Table 2.6
J_VALUES = {
    "JPCP + Dowel + Tied Shoulder": 2.8,
    "JPCP + Dowel Bar (AC Shoulder)": 3.2,
    "JPCP ไม่มี Dowel Bar": 3.8,
    "CRCP + Tied Shoulder": 2.5,
    "CRCP (AC Shoulder)": 2.9
}

# ค่า Drainage Coefficient (Cd) มาตรฐาน
CD_DEFAULT = 1.0

# ============================================================
# ส่วนที่ 2: ฟังก์ชันการคำนวณ
# ============================================================

def convert_cube_to_cylinder(fc_cube_ksc: float) -> float:
    """
    แปลงกำลังอัดคอนกรีตจาก Cube เป็น Cylinder
    fc_cylinder ≈ 0.833 × fc_cube (โดยประมาณ)
    
    Parameters:
        fc_cube_ksc: กำลังอัดคอนกรีต Cube (ksc)
    
    Returns:
        กำลังอัดคอนกรีต Cylinder (ksc)
    """
    return 0.8 * fc_cube_ksc


def calculate_concrete_modulus(fc_cylinder_ksc: float) -> float:
    """
    คำนวณ Modulus of Elasticity ของคอนกรีต (Ec)
    ตามสูตร ACI: Ec = 57,000 × √(f'c) (psi)
    
    Parameters:
        fc_cylinder_ksc: กำลังอัดคอนกรีต Cylinder (ksc)
    
    Returns:
        Ec ในหน่วย psi
    """
    # แปลง ksc เป็น psi (1 ksc = 14.223 psi)
    fc_psi = fc_cylinder_ksc * 14.223
    
    # คำนวณ Ec ตาม ACI 318
    ec_psi = 57000 * math.sqrt(fc_psi)
    
    return ec_psi


def estimate_modulus_of_rupture(fc_cylinder_ksc: float) -> float:
    """
    ประมาณค่า Modulus of Rupture (Sc) จากกำลังอัดคอนกรีต
    ตามสูตร: Sc = (7.5 ถึง 12) × √(f'c) (ACI 318, หน่วย psi)
    
    Parameters:
        fc_cylinder_ksc: กำลังอัดคอนกรีต Cylinder (ksc)
    
    Returns:
        Sc ในหน่วย psi (ใช้ค่า 10 × √f'c)
    """
    # แปลง ksc เป็น psi
    fc_psi = fc_cylinder_ksc * 14.223
    
    # ใช้สูตร: Sc = 10 × √f'c (ค่าเหมาะสมสำหรับคอนกรีตถนน)
    sc_psi = 10.0 * math.sqrt(fc_psi)
    
    return sc_psi


def get_zr_value(reliability: float) -> float:
    """
    หาค่า ZR จากตาราง AASHTO ตามระดับความเชื่อมั่น
    
    Parameters:
        reliability: ระดับความเชื่อมั่น (%)
    
    Returns:
        ค่า ZR (Standard Normal Deviate)
    """
    return ZR_TABLE.get(int(reliability), -1.282)


def calculate_aashto_rigid_w18(
    d_inch: float,
    delta_psi: float,
    pt: float,
    zr: float,
    so: float,
    sc_psi: float,
    cd: float,
    j: float,
    ec_psi: float,
    k_pci: float
) -> tuple:
    """
    คำนวณ ESAL (W18) ที่รองรับได้ตามสมการ AASHTO 1993 สำหรับ Rigid Pavement
    
    สมการ AASHTO 1993:
    log10(W18) = ZR × So + 7.35 × log10(D+1) - 0.06 
                 + log10(ΔPSI/(4.5-1.5)) / (1 + 1.624×10^7 / (D+1)^8.46)
                 + (4.22 - 0.32×Pt) × log10[(Sc×Cd×(D^0.75-1.132)) / (215.63×J×(D^0.75 - 18.42/(Ec/k)^0.25))]
    
    Parameters:
        d_inch: ความหนาแผ่นพื้นคอนกรีต (นิ้ว)
        delta_psi: การสูญเสียค่า Serviceability (ΔPSI = 4.5 - Pt)
        pt: Terminal Serviceability
        zr: Standard Normal Deviate
        so: Overall Standard Deviation
        sc_psi: Modulus of Rupture (psi)
        cd: Drainage Coefficient
        j: Load Transfer Coefficient
        ec_psi: Modulus of Elasticity ของคอนกรีต (psi)
        k_pci: Effective Modulus of Subgrade Reaction (pci)
    
    Returns:
        tuple: (log10_w18, w18)
    """
    
    # พจน์ที่ 1: ZR × So
    term1 = zr * so
    
    # พจน์ที่ 2: 7.35 × log10(D+1) - 0.06
    term2 = 7.35 * math.log10(d_inch + 1) - 0.06
    
    # พจน์ที่ 3: การสูญเสีย Serviceability
    # log10(ΔPSI/(4.5-1.5)) / (1 + 1.624×10^7 / (D+1)^8.46)
    numerator3 = math.log10(delta_psi / (4.5 - 1.5))
    denominator3 = 1 + (1.624e7 / ((d_inch + 1) ** 8.46))
    term3 = numerator3 / denominator3
    
    # พจน์ที่ 4: กำลังของคอนกรีตและฐานราก
    # (4.22 - 0.32×Pt) × log10[(Sc×Cd×(D^0.75-1.132)) / (215.63×J×(D^0.75 - 18.42/(Ec/k)^0.25))]
    
    # คำนวณ D^0.75
    d_power = d_inch ** 0.75
    
    # คำนวณตัวเศษ: Sc × Cd × (D^0.75 - 1.132)
    numerator4 = sc_psi * cd * (d_power - 1.132)
    
    # คำนวณตัวส่วน: 215.63 × J × (D^0.75 - 18.42/(Ec/k)^0.25)
    ec_k_ratio = ec_psi / k_pci
    denominator4 = 215.63 * j * (d_power - 18.42 / (ec_k_ratio ** 0.25))
    
    # ตรวจสอบว่าค่าต้องเป็นบวก
    if numerator4 <= 0 or denominator4 <= 0:
        return (float('-inf'), 0)
    
    inner_term = numerator4 / denominator4
    
    if inner_term <= 0:
        return (float('-inf'), 0)
    
    term4 = (4.22 - 0.32 * pt) * math.log10(inner_term)
    
    # รวมทุกพจน์
    log10_w18 = term1 + term2 + term3 + term4
    
    # คำนวณ W18
    w18 = 10 ** log10_w18
    
    return (log10_w18, w18)


def check_design(w18_required: float, w18_capacity: float) -> tuple:
    """
    ตรวจสอบว่าความหนาที่กำหนดรองรับ ESAL ได้หรือไม่
    
    Parameters:
        w18_required: ESAL ที่ต้องการรองรับ
        w18_capacity: ESAL ที่รองรับได้
    
    Returns:
        tuple: (ผลการตรวจสอบ (bool), อัตราส่วน)
    """
    ratio = w18_capacity / w18_required if w18_required > 0 else float('inf')
    passed = w18_capacity >= w18_required
    return (passed, ratio)


# ============================================================
# ส่วนที่ 3: ฟังก์ชันสร้างรายงาน Word
# ============================================================

def create_word_report(
    pavement_type: str,
    inputs: dict,
    calculated_values: dict,
    comparison_results: list,
    selected_d: float,
    main_result: tuple,
    layers_data: list = None
) -> BytesIO:
    """
    สร้างรายงานการคำนวณในรูปแบบไฟล์ Word (.docx)
    ใช้ python-docx library
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        st.error("กรุณาติดตั้ง python-docx: pip install python-docx")
        return None
    
    # สร้างเอกสารใหม่
    doc = Document()
    
    # ตั้งค่าฟอนต์ภาษาไทย
    style = doc.styles['Normal']
    font = style.font
    font.name = 'TH Sarabun New'
    font.size = Pt(14)
    
    # หัวข้อเอกสาร
    title = doc.add_heading('รายการคำนวณออกแบบความหนาถนนคอนกรีต', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('ตามวิธี AASHTO 1993')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ข้อมูลทั่วไป
    doc.add_heading('1. ข้อมูลทั่วไป', level=1)
    doc.add_paragraph(f'ประเภทถนน: {pavement_type}')
    doc.add_paragraph(f'วันที่คำนวณ: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    
    # ตารางชั้นโครงสร้างทาง
    if layers_data and len(layers_data) > 0:
        doc.add_heading('2. ชั้นโครงสร้างทาง (Pavement Layers)', level=1)
        
        table_layers = doc.add_table(rows=1, cols=4)
        table_layers.style = 'Table Grid'
        hdr_layers = table_layers.rows[0].cells
        hdr_layers[0].text = 'ลำดับ'
        hdr_layers[1].text = 'ชนิดวัสดุ'
        hdr_layers[2].text = 'ความหนา (ซม.)'
        hdr_layers[3].text = 'Modulus E (MPa)'
        
        for i, layer in enumerate(layers_data):
            row_cells = table_layers.add_row().cells
            row_cells[0].text = str(i + 1)
            row_cells[1].text = layer.get('name', f'Layer {i+1}')
            row_cells[2].text = f"{layer.get('thickness_cm', 0)}"
            row_cells[3].text = f"{layer.get('E_MPa', 0):,}"
        
        doc.add_paragraph('')  # เว้นบรรทัด
    
    # ข้อมูลนำเข้า
    doc.add_heading('3. ข้อมูลนำเข้า (Input Parameters)', level=1)
    
    table1 = doc.add_table(rows=1, cols=4)
    table1.style = 'Table Grid'
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = 'พารามิเตอร์'
    hdr_cells[1].text = 'สัญลักษณ์'
    hdr_cells[2].text = 'ค่า'
    hdr_cells[3].text = 'หน่วย'
    
    input_data = [
        ('ESAL ออกแบบ', 'W₁₈', f"{inputs['w18_design']:,.0f}", 'ESALs'),
        ('Terminal Serviceability', 'Pt', f"{inputs['pt']:.1f}", '-'),
        ('Reliability', 'R', f"{inputs['reliability']:.0f}", '%'),
        ('Standard Deviation', 'So', f"{inputs['so']:.2f}", '-'),
        ('Modulus of Subgrade Reaction', 'k_eff', f"{inputs['k_eff']:,.0f}", 'pci'),
        ('Loss of Support', 'LS', f"{inputs.get('ls', 1.0):.1f}", '-'),
        ('กำลังคอนกรีต', "f'c", f"{inputs['fc_cube']:.0f} Cube ({int(inputs['fc_cube']*0.8)} Cyl.)", 'ksc'),
        ('Modulus of Rupture', 'Sc', f"{inputs['sc']:.0f}", 'psi'),
        ('Load Transfer Coefficient', 'J', f"{inputs['j']:.1f}", '-'),
        ('Drainage Coefficient', 'Cd', f"{inputs['cd']:.1f}", '-'),
    ]
    
    for param, symbol, value, unit in input_data:
        row_cells = table1.add_row().cells
        row_cells[0].text = param
        row_cells[1].text = symbol
        row_cells[2].text = value
        row_cells[3].text = unit
    
    # ค่าที่คำนวณได้
    doc.add_heading('4. ค่าที่คำนวณได้ (Calculated Values)', level=1)
    
    table2 = doc.add_table(rows=1, cols=4)
    table2.style = 'Table Grid'
    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = 'พารามิเตอร์'
    hdr_cells2[1].text = 'สัญลักษณ์'
    hdr_cells2[2].text = 'ค่า'
    hdr_cells2[3].text = 'หน่วย'
    
    calc_data = [
        ('Modulus of Elasticity', 'Ec', f"{calculated_values['ec']:,.0f}", 'psi'),
        ('Standard Normal Deviate', 'ZR', f"{calculated_values['zr']:.3f}", '-'),
        ('การสูญเสีย Serviceability', 'ΔPSI', f"{calculated_values['delta_psi']:.1f}", '-'),
    ]
    
    for param, symbol, value, unit in calc_data:
        row_cells = table2.add_row().cells
        row_cells[0].text = param
        row_cells[1].text = symbol
        row_cells[2].text = value
        row_cells[3].text = unit
    
    # สมการ AASHTO 1993
    doc.add_heading('5. สมการออกแบบ AASHTO 1993', level=1)
    
    equation_text = """
    log₁₀(W₁₈) = ZR × So + 7.35 × log₁₀(D+1) - 0.06 
                 + log₁₀(ΔPSI/(4.5-1.5)) / (1 + 1.624×10⁷/(D+1)^8.46)
                 + (4.22 - 0.32×Pt) × log₁₀[(Sc×Cd×(D^0.75-1.132))/(215.63×J×(D^0.75 - 18.42/(Ec/k)^0.25))]
    """
    doc.add_paragraph(equation_text)
    
    # ผลการเปรียบเทียบ
    doc.add_heading('6. ผลการเปรียบเทียบความหนาต่างๆ', level=1)
    
    table3 = doc.add_table(rows=1, cols=5)
    table3.style = 'Table Grid'
    hdr_cells3 = table3.rows[0].cells
    hdr_cells3[0].text = 'D (นิ้ว)'
    hdr_cells3[1].text = 'log₁₀(W₁₈)'
    hdr_cells3[2].text = 'W₁₈ รองรับได้'
    hdr_cells3[3].text = 'อัตราส่วน'
    hdr_cells3[4].text = 'ผลการตรวจสอบ'
    
    for result in comparison_results:
        row_cells = table3.add_row().cells
        row_cells[0].text = f"{result['d']:.0f}"
        row_cells[1].text = f"{result['log_w18']:.4f}"
        row_cells[2].text = f"{result['w18']:,.0f}"
        row_cells[3].text = f"{result['ratio']:.2f}"
        row_cells[4].text = "ผ่าน ✓" if result['passed'] else "ไม่ผ่าน ✗"
    
    # สรุปผล
    doc.add_heading('7. สรุปผลการออกแบบ', level=1)
    
    passed, ratio = main_result
    status = "ผ่านเกณฑ์ ✓" if passed else "ไม่ผ่านเกณฑ์ ✗"
    
    summary = f"""
    ความหนาที่เลือก: {selected_d:.0f} นิ้ว ({selected_d * 2.54:.1f} ซม.)
    ESAL ที่ต้องการ: {inputs['w18_design']:,.0f} ESALs
    ESAL ที่รองรับได้: {[r for r in comparison_results if r['d'] == selected_d][0]['w18'] if any(r['d'] == selected_d for r in comparison_results) else 'N/A':,.0f} ESALs
    อัตราส่วน: {ratio:.2f}
    ผลการตรวจสอบ: {status}
    """
    doc.add_paragraph(summary)
    
    # หมายเหตุ
    doc.add_heading('8. หมายเหตุ', level=1)
    notes = """
    - การคำนวณนี้ใช้หลักการตามคู่มือ AASHTO Guide for Design of Pavement Structures (1993)
    - สมการ: log₁₀(W₁₈) รวม term (D^0.75 - 1.132) ในตัวเศษ
    - ค่า J สำหรับ JPCP + Dowel + Tied Shoulder = 2.7, JPCP + Dowel (AC Shoulder) = 3.2
    - การแปลงกำลังคอนกรีต: f'c (cylinder) ≈ 0.8 × f'c (cube)
    - Ec = 57,000 × √f'c (psi) ตาม ACI 318
    - Sc ≈ 10 × √f'c (psi)
    """
    doc.add_paragraph(notes)
    
    # บันทึกไฟล์ลง BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


# ============================================================
# ส่วนที่ 4: Streamlit UI
# ============================================================

def main():
    # ตั้งค่าหน้าเว็บ
    st.set_page_config(
        page_title="AASHTO 1993 Rigid Pavement Design",
        page_icon="🛣️",
        layout="wide"
    )
    
    # หัวข้อหลัก
    st.title("🛣️ การออกแบบความหนาถนนคอนกรีต")
    st.subheader("ตามวิธี AASHTO 1993 (Rigid Pavement Design)")
    
    st.markdown("---")
    
    # แบ่งคอลัมน์
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.header("📥 ข้อมูลนำเข้า (Input)")
        
        # เลือกประเภทถนน
        pavement_type = st.selectbox(
            "ประเภทถนนคอนกรีต",
            options=list(J_VALUES.keys()),
            index=0,
            help="JPCP = Jointed Plain Concrete Pavement, CRCP = Continuously Reinforced Concrete Pavement"
        )
        
        st.markdown("---")
        
        # ชั้นโครงสร้างทาง (Pavement Layers)
        st.subheader("🔶 ชั้นโครงสร้างทาง (Pavement Layers)")
        
        # จำนวนชั้นวัสดุ
        num_layers = st.slider(
            "จำนวนชั้นวัสดุใต้แผ่นคอนกรีต",
            min_value=1,
            max_value=6,
            value=5,
            help="เลือกจำนวนชั้นวัสดุ 1-6 ชั้น"
        )
        
        # ค่า Default สำหรับแต่ละชั้น
        default_layers = [
            {"name": "Asphalt Treated Base", "thickness_cm": 5, "E_MPa": 2500},
            {"name": "Cement Treated Base", "thickness_cm": 20, "E_MPa": 1200},
            {"name": "Crusite Base", "thickness_cm": 15, "E_MPa": 150},
            {"name": "Granular Subbase", "thickness_cm": 25, "E_MPa": 80},
            {"name": "Selected Material", "thickness_cm": 30, "E_MPa": 50},
            {"name": "Subgrade", "thickness_cm": 0, "E_MPa": 30},
        ]
        
        # เก็บข้อมูลชั้นวัสดุ
        layers_data = []
        
        with st.expander("📊 ตารางค่า Modulus แนะนำ", expanded=False):
            st.markdown("""
            | ประเภทวัสดุ | E (MPa) |
            |------------|---------|
            | Asphalt Treated Base | 1,500 - 3,000 |
            | Cement Treated Base | 1,000 - 2,000 |
            | Crushed Stone Base | 100 - 300 |
            | Granular Subbase | 50 - 150 |
            | Selected Material | 30 - 80 |
            | Subgrade (ดินเดิม) | 20 - 50 |
            """)
        
        for i in range(num_layers):
            st.markdown(f"**ชั้นที่ {i+1}**")
            col_a, col_b, col_c = st.columns([2, 1, 1])
            
            with col_a:
                layer_name = st.text_input(
                    f"ชื่อวัสดุ",
                    value=default_layers[i]["name"] if i < len(default_layers) else f"Layer {i+1}",
                    key=f"layer_name_{i}"
                )
            
            with col_b:
                layer_thickness = st.number_input(
                    f"ความหนา (ซม.)",
                    min_value=0,
                    max_value=100,
                    value=default_layers[i]["thickness_cm"] if i < len(default_layers) else 20,
                    key=f"layer_thick_{i}"
                )
            
            with col_c:
                layer_modulus = st.number_input(
                    f"E (MPa)",
                    min_value=10,
                    max_value=10000,
                    value=default_layers[i]["E_MPa"] if i < len(default_layers) else 100,
                    key=f"layer_E_{i}"
                )
            
            layers_data.append({
                "name": layer_name,
                "thickness_cm": layer_thickness,
                "E_MPa": layer_modulus
            })
        
        st.markdown("---")
        
        # 1. ESAL ที่ต้องการรองรับ
        st.subheader("1️⃣ ปริมาณจราจร")
        
        # แสดงตัวช่วยประมาณ ESAL
        with st.expander("📊 ตัวช่วยประมาณ ESAL ตามประเภทถนน"):
            st.markdown("""
            | ประเภทถนน | ESAL (20 ปี) |
            |-----------|--------------|
            | ถนนในหมู่บ้าน | 50,000 - 200,000 |
            | ถนนเทศบาล | 200,000 - 500,000 |
            | ถนน อบจ. / ทางหลวงชนบท | 500,000 - 2,000,000 |
            | ทางหลวงแผ่นดิน (2 ช่องจราจร) | 2,000,000 - 10,000,000 |
            | ทางหลวงแผ่นดิน (4 ช่องจราจร) | 10,000,000 - 50,000,000 |
            """)
        
        w18_design = st.number_input(
            "ESAL ที่ต้องการรองรับ (W₁₈)",
            min_value=10_000,
            max_value=500_000_000,
            value=500_000,
            step=100_000,
            format="%d",
            help="จำนวน Equivalent Single Axle Load (18 kip) ตลอดอายุการใช้งาน"
        )
        
        st.markdown("---")
        
        # 2. Serviceability
        st.subheader("2️⃣ Serviceability")
        pt = st.slider(
            "Terminal Serviceability (Pt)",
            min_value=1.5,
            max_value=3.0,
            value=2.0,
            step=0.1,
            help="ค่า Serviceability ที่ยอมรับได้ต่ำสุด (มาตรฐาน = 2.0)"
        )
        
        # คำนวณ ΔPSI
        delta_psi = 4.5 - pt
        st.info(f"ΔPSI = 4.5 - {pt:.1f} = **{delta_psi:.1f}**")
        
        st.markdown("---")
        
        # 3. Reliability
        st.subheader("3️⃣ ความเชื่อมั่นในการออกแบบ")
        reliability = st.select_slider(
            "Reliability (R)",
            options=[80, 85, 90, 95],
            value=90,
            help="ระดับความเชื่อมั่นในการออกแบบ (%)"
        )
        
        # หาค่า ZR
        zr = get_zr_value(reliability)
        st.info(f"ZR = **{zr:.3f}** (จากตาราง AASHTO)")
        
        # Standard Deviation
        so = st.number_input(
            "Overall Standard Deviation (So)",
            min_value=0.30,
            max_value=0.45,
            value=0.35,
            step=0.01,
            format="%.2f",
            help="ค่าเบี่ยงเบนมาตรฐานรวม (มาตรฐาน = 0.35 สำหรับ Rigid Pavement)"
        )
        
        st.markdown("---")
        
        # 4. คุณสมบัติดินฐานราก
        st.subheader("4️⃣ คุณสมบัติดินฐานราก")
        
        with st.expander("📊 ตารางประมาณค่า k จาก CBR"):
            st.markdown("""
            | CBR (%) | k (pci) | คำอธิบาย |
            |---------|---------|----------|
            | 2-3 | 75-100 | ดินเหนียวอ่อน |
            | 4-5 | 100-130 | ดินเหนียวแข็ง |
            | 6-10 | 130-170 | ดินทรายปนดินเหนียว |
            | 10-20 | 170-230 | ดินทรายอัดแน่น |
            | 20-50 | 230-350 | หินคลุก/ลูกรัง |
            | > 50 | 350-500+ | ชั้น Base คุณภาพดี |
            
            **หมายเหตุ:** ค่า k_eff รวมผลของชั้น Subbase แล้ว
            """)
        
        k_eff = st.number_input(
            "Effective Modulus of Subgrade Reaction (k_eff)",
            min_value=50,
            max_value=1000,
            value=200,
            step=25,
            format="%d",
            help="ค่า k จากการทดสอบ Plate Bearing Test หรือประมาณจาก CBR (หน่วย: pci)"
        )
        
        # Loss of Support (LS)
        st.markdown("**Loss of Support (LS)**")
        
        with st.expander("📊 ตารางค่า Loss of Support แนะนำ (AASHTO 1993)"):
            st.markdown("""
            | ประเภทวัสดุ | Loss of Support (LS) |
            |------------|---------------------|
            | Cement Treated Granular Base | 0.0 - 1.0 |
            | Cement Aggregate Mixtures | 0.0 - 1.0 |
            | Asphalt Treated Base | 0.0 - 1.0 |
            | Bituminous Stabilized Mixtures | 0.0 - 1.0 |
            | Lime Stabilized | 1.0 - 3.0 |
            | Unbound Granular Materials | 1.0 - 3.0 |
            | Fine Grained or Natural Subgrade | 2.0 - 3.0 |
            
            **หมายเหตุ:** ค่า LS ใช้ปรับลดค่า k_eff เพื่อคำนึงถึงการสูญเสียการรองรับจากการกัดเซาะ
            """)
        
        ls_value = st.number_input(
            "ค่า Loss of Support (LS)",
            min_value=0.0,
            max_value=3.0,
            value=1.0,
            step=0.5,
            format="%.1f",
            help="ค่า LS สำหรับปรับลด k_eff (0.0-3.0)"
        )
        
        st.markdown("---")
        
        # 5. คุณสมบัติคอนกรีต
        st.subheader("5️⃣ คุณสมบัติคอนกรีต")
        
        fc_cube = st.number_input(
            "กำลังอัดคอนกรีต (Cube) - f'c",
            min_value=200,
            max_value=600,
            value=350,
            step=10,
            format="%d",
            help="กำลังอัดคอนกรีตที่ 28 วัน ทดสอบด้วย Cube 15×15×15 ซม. (หน่วย: ksc)"
        )
        
        # แปลง Cube เป็น Cylinder
        fc_cylinder = convert_cube_to_cylinder(fc_cube)
        st.info(f"f'c (Cylinder) = 0.8 × {fc_cube} = **{fc_cylinder:.0f} ksc**")
        
        # คำนวณ Ec
        ec = calculate_concrete_modulus(fc_cylinder)
        st.info(f"Ec = 57,000 × √({fc_cylinder * 14.223:.0f}) = **{ec:,.0f} psi**")
        
        # Modulus of Rupture
        st.markdown("**Modulus of Rupture (Sc)**")
        
        # คำนวณค่า Sc อัตโนมัติ
        sc_auto = estimate_modulus_of_rupture(fc_cylinder)
        st.info(f"ค่าประมาณ: Sc = 10 × √({fc_cylinder * 14.223:.0f}) = **{sc_auto:.0f} psi**")
        
        # ให้ผู้ใช้ป้อนค่าที่ต้องการใช้
        sc = st.number_input(
            "ค่า Sc ที่ใช้ในการคำนวณ (psi)",
            min_value=400,
            max_value=1000,
            value=int(round(sc_auto)),
            step=10,
            format="%d",
            help="ค่าเริ่มต้นคำนวณจาก 10×√f'c สามารถแก้ไขได้ตามผลทดสอบจริง"
        )
        
        st.markdown("---")
        
        # 6. Load Transfer และ Drainage
        st.subheader("6️⃣ Load Transfer และ Drainage")
        
        # แสดงค่า J อัตโนมัติตามประเภทถนน
        j_auto = J_VALUES[pavement_type]
        st.info(f"ค่าแนะนำสำหรับ {pavement_type}: **J = {j_auto}**")
        
        # ตารางอ้างอิงค่า J
        with st.expander("📊 ตารางค่า Load Transfer Coefficient (J)"):
            st.markdown("""
            | ประเภทถนน | J (Tied Shoulder) | J (AC Shoulder) |
            |-----------|-------------------|-----------------|
            | JPCP + Dowel Bar | 2.7 | 3.2 |
            | JPCP ไม่มี Dowel | 3.2 | 3.8-4.4 |
            | CRCP | 2.3 | 2.9 |
            
            **หมายเหตุ:** ค่า J ต่ำ = การถ่ายแรงดี = รองรับ ESAL ได้มากขึ้น
            """)
        
        # ให้ผู้ใช้ป้อนค่าที่ต้องการใช้
        j_value = st.number_input(
            "ค่า J ที่ใช้ในการคำนวณ",
            min_value=2.0,
            max_value=4.5,
            value=j_auto,
            step=0.1,
            format="%.1f",
            help="ค่าเริ่มต้นตามประเภทถนนที่เลือก สามารถแก้ไขได้"
        )
        
        cd = st.number_input(
            "Drainage Coefficient (Cd)",
            min_value=0.7,
            max_value=1.3,
            value=1.0,
            step=0.05,
            format="%.2f",
            help="สัมประสิทธิ์การระบายน้ำ (1.0 = การระบายน้ำปานกลาง)"
        )
        
        st.markdown("---")
        
        # 7. ความหนาคอนกรีต
        st.subheader("7️⃣ ความหนาคอนกรีตที่ต้องการตรวจสอบ")
        d_selected = st.slider(
            "ความหนาคอนกรีต D (นิ้ว)",
            min_value=8,
            max_value=16,
            value=12,
            step=1,
            help="ความหนาแผ่นพื้นคอนกรีต"
        )
        st.info(f"D = {d_selected} นิ้ว = **{d_selected * 2.54:.1f} ซม.**")
    
    # ============================================================
    # ส่วนแสดงผลการคำนวณ
    # ============================================================
    
    with col2:
        st.header("📊 ผลการคำนวณ (Output)")
        
        # เก็บผลการคำนวณสำหรับความหนาต่างๆ
        comparison_results = []
        thicknesses = [8, 9, 10, 11, 12, 13, 14, 15, 16]
        
        # คำนวณสำหรับแต่ละความหนา
        st.subheader("📋 ตารางเปรียบเทียบความหนาต่างๆ")
        
        # สร้างตาราง
        table_data = []
        for d in thicknesses:
            log_w18, w18_capacity = calculate_aashto_rigid_w18(
                d_inch=d,
                delta_psi=delta_psi,
                pt=pt,
                zr=zr,
                so=so,
                sc_psi=sc,
                cd=cd,
                j=j_value,
                ec_psi=ec,
                k_pci=k_eff
            )
            passed, ratio = check_design(w18_design, w18_capacity)
            
            comparison_results.append({
                'd': d,
                'log_w18': log_w18,
                'w18': w18_capacity,
                'passed': passed,
                'ratio': ratio
            })
            
            table_data.append({
                'D (นิ้ว)': d,
                'D (ซม.)': f"{d * 2.54:.1f}",
                'log₁₀(W₁₈)': f"{log_w18:.4f}",
                'W₁₈ รองรับได้': f"{w18_capacity:,.0f}",
                'อัตราส่วน': f"{ratio:.2f}",
                'ผล': "✅ ผ่าน" if passed else "❌ ไม่ผ่าน"
            })
        
        # แสดงตาราง
        import pandas as pd
        df = pd.DataFrame(table_data)
        
        # จัดรูปแบบตาราง
        st.dataframe(
            df,
            use_container_width=True,
            hide_index=True
        )
        
        st.markdown("---")
        
        # ผลการคำนวณสำหรับความหนาที่เลือก
        st.subheader(f"🎯 ผลการตรวจสอบ D = {d_selected} นิ้ว")
        
        log_w18_selected, w18_selected = calculate_aashto_rigid_w18(
            d_inch=d_selected,
            delta_psi=delta_psi,
            pt=pt,
            zr=zr,
            so=so,
            sc_psi=sc,
            cd=cd,
            j=j_value,
            ec_psi=ec,
            k_pci=k_eff
        )
        passed_selected, ratio_selected = check_design(w18_design, w18_selected)
        
        # แสดงผลด้วยสี
        col_a, col_b = st.columns(2)
        
        with col_a:
            st.metric(
                label="log₁₀(W₁₈)",
                value=f"{log_w18_selected:.4f}"
            )
            st.metric(
                label="W₁₈ รองรับได้",
                value=f"{w18_selected:,.0f}",
                delta=f"{w18_selected - w18_design:+,.0f}"
            )
        
        with col_b:
            st.metric(
                label="W₁₈ ที่ต้องการ",
                value=f"{w18_design:,.0f}"
            )
            st.metric(
                label="อัตราส่วน (Capacity/Required)",
                value=f"{ratio_selected:.2f}"
            )
        
        # แสดงผลผ่าน/ไม่ผ่าน
        if passed_selected:
            st.success(f"""
            ✅ **ผ่านเกณฑ์การออกแบบ**
            
            ความหนา D = {d_selected} นิ้ว ({d_selected * 2.54:.1f} ซม.) 
            สามารถรองรับ ESAL ได้ {w18_selected:,.0f} ESALs
            ซึ่งมากกว่า ESAL ที่ต้องการ {w18_design:,.0f} ESALs
            
            อัตราส่วน = {ratio_selected:.2f} (≥ 1.00)
            """)
        else:
            st.error(f"""
            ❌ **ไม่ผ่านเกณฑ์การออกแบบ**
            
            ความหนา D = {d_selected} นิ้ว ({d_selected * 2.54:.1f} ซม.) 
            รองรับ ESAL ได้เพียง {w18_selected:,.0f} ESALs
            ซึ่งน้อยกว่า ESAL ที่ต้องการ {w18_design:,.0f} ESALs
            
            อัตราส่วน = {ratio_selected:.2f} (< 1.00)
            
            **กรุณาเพิ่มความหนาคอนกรีต หรือปรับปรุงคุณสมบัติวัสดุ**
            """)
        
        st.markdown("---")
        
        # แสดงสมการที่ใช้
        st.subheader("📝 สมการ AASHTO 1993")
        
        st.latex(r'''
        \log_{10}(W_{18}) = Z_R \times S_o + 7.35 \times \log_{10}(D+1) - 0.06
        ''')
        
        st.latex(r'''
        + \frac{\log_{10}\left(\frac{\Delta PSI}{4.5-1.5}\right)}{1 + \frac{1.624 \times 10^7}{(D+1)^{8.46}}}
        ''')
        
        st.latex(r'''
        + (4.22 - 0.32 \times P_t) \times \log_{10}\left[\frac{S_c \times C_d \times (D^{0.75} - 1.132)}{215.63 \times J \times \left(D^{0.75} - \frac{18.42}{(E_c/k)^{0.25}}\right)}\right]
        ''')
        
        st.markdown("---")
        
        # ส่งออกรายงาน Word
        st.subheader("📄 ส่งออกรายงาน")
        
        # เตรียมข้อมูลสำหรับรายงาน
        inputs_dict = {
            'w18_design': w18_design,
            'pt': pt,
            'reliability': reliability,
            'so': so,
            'k_eff': k_eff,
            'ls': ls_value,
            'fc_cube': fc_cube,
            'sc': sc,
            'j': j_value,
            'cd': cd
        }
        
        calculated_dict = {
            'fc_cylinder': fc_cylinder,
            'ec': ec,
            'zr': zr,
            'delta_psi': delta_psi
        }
        
        # สร้างปุ่มดาวน์โหลด
        if st.button("📥 สร้างรายงาน Word", type="primary"):
            with st.spinner("กำลังสร้างรายงาน..."):
                try:
                    buffer = create_word_report(
                        pavement_type=pavement_type,
                        inputs=inputs_dict,
                        calculated_values=calculated_dict,
                        comparison_results=comparison_results,
                        selected_d=d_selected,
                        main_result=(passed_selected, ratio_selected),
                        layers_data=layers_data
                    )
                    
                    if buffer:
                        st.download_button(
                            label="⬇️ ดาวน์โหลดรายงาน (.docx)",
                            data=buffer,
                            file_name=f"AASHTO_Rigid_Pavement_Design_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        st.success("สร้างรายงานสำเร็จ!")
                except Exception as e:
                    st.error(f"เกิดข้อผิดพลาด: {str(e)}")
                    st.info("กรุณาติดตั้ง python-docx: `pip install python-docx`")
    
    # ============================================================
    # ส่วนอ้างอิง
    # ============================================================
    
    st.markdown("---")
    st.header("📚 อ้างอิง")
    
    st.markdown("""
    **เอกสารอ้างอิง:**
    1. AASHTO (1993). *AASHTO Guide for Design of Pavement Structures*. American Association of State Highway and Transportation Officials.
    2. Huang, Y.H. (2004). *Pavement Analysis and Design*. Pearson Prentice Hall.
    3. ACI 318-19 (2019). *Building Code Requirements for Structural Concrete*. American Concrete Institute.
    
    **หมายเหตุ:**
    - โปรแกรมนี้พัฒนาเพื่อใช้ในการเรียนการสอน
    - การออกแบบจริงควรพิจารณาปัจจัยอื่นๆ ร่วมด้วย เช่น สภาพแวดล้อม การก่อสร้าง และการบำรุงรักษา
    """)
    
    # Footer
    st.markdown("---")
    st.caption("พัฒนาโดย: ภาควิชาครุศาสตร์โยธา มจพ. | AASHTO 1993 Rigid Pavement Design Tool")


if __name__ == "__main__":
    main()
