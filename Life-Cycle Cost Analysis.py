#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
================================================================================
โปรแกรมวิเคราะห์ต้นทุนตลอดอายุการใช้งานผิวทาง (LCCA) - เวอร์ชัน 2.1
Pavement Life-Cycle Cost Analysis Program
================================================================================
พัฒนาสำหรับการเรียนการสอนและงานวิจัยด้านวิศวกรรมทาง
ภาควิชาครุศาสตร์โยธา มหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ

คุณสมบัติเวอร์ชัน 2.1:
- 📤 Upload Excel เพื่อกรอกข้อมูลต้นทุนก่อสร้าง
- 📥 ดาวน์โหลด Template Excel
- 🔄 เปิด/ปิดการคำนวณมูลค่าซาก
- ⚡ Fast upload method (ไม่ต้องรอ reload)
- แก้ไขต้นทุนก่อสร้างได้เอง
- กำหนดพื้นที่โครงการได้เอง
- เพิ่มผิวทาง JRCP (Jointed Reinforced Concrete Pavement)
- แก้ไขแผนบำรุงรักษาและฟื้นฟูสภาพได้

ประเภทผิวทาง:
1. Flexible Pavement (ผิวทางยืดหยุ่น/แอสฟัลต์)
2. JPCP - Jointed Plain Concrete Pavement (คอนกรีตไม่เสริมเหล็ก)
3. JRCP - Jointed Reinforced Concrete Pavement (คอนกรีตเสริมเหล็ก)
4. CRCP - Continuously Reinforced Concrete Pavement (คอนกรีตเสริมเหล็กต่อเนื่อง)
================================================================================
"""

import streamlit as st
import numpy as np
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
from typing import Dict, List, Tuple, Optional
from dataclasses import dataclass, field
import json
import io
from datetime import datetime
import hashlib

# สำหรับส่งออก Word
try:
    from docx import Document as WordDocument
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ตั้งค่าหน้าเว็บ
st.set_page_config(
    page_title="โปรแกรมวิเคราะห์ LCCA ผิวทาง v2.1",
    page_icon="🛣️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# =============================================================================
# ส่วนที่ 1: โครงสร้างข้อมูล (Data Structures)
# =============================================================================

@dataclass
class กิจกรรมบำรุงรักษา:
    """โครงสร้างข้อมูลกิจกรรมบำรุงรักษา"""
    ชื่อกิจกรรม: str
    ต้นทุนต่อหน่วย: float  # บาท/ตร.ม.
    ปีเริ่มต้น: int
    ความถี่: int = 0  # 0 = ครั้งเดียว


@dataclass
class กิจกรรมฟื้นฟูสภาพ:
    """โครงสร้างข้อมูลกิจกรรมฟื้นฟูสภาพ"""
    ชื่อกิจกรรม: str
    ต้นทุนต่อหน่วย: float  # บาท/ตร.ม.
    ปีดำเนินการ: int


@dataclass
class ทางเลือกผิวทาง:
    """โครงสร้างข้อมูลทางเลือกผิวทาง"""
    ชื่อ: str
    ประเภท: str
    ต้นทุนก่อสร้าง: float  # บาท/ตร.ม.
    แผนบำรุงรักษา: List[กิจกรรมบำรุงรักษา]
    แผนฟื้นฟูสภาพ: List[กิจกรรมฟื้นฟูสภาพ]
    ร้อยละมูลค่าซาก: float = 20.0
    พื้นที่: float = 1000.0  # ตร.ม.
    ความหนา: float = 0.0  # ซม. (0 = ไม่ระบุ)
    เปิดใช้งาน: bool = True


# =============================================================================
# ส่วนที่ 2: ฟังก์ชันคำนวณหลัก (Core Calculation Functions)
# =============================================================================

def คำนวณมูลค่าปัจจุบัน(ต้นทุน: float, ปี: int, อัตราคิดลด: float) -> float:
    """
    คำนวณมูลค่าปัจจุบัน (Present Worth)
    สูตร: PW = FV × (1 + i)^(-n)
    """
    if ปี < 0 or อัตราคิดลด < 0:
        return 0.0
    pwf = (1 + อัตราคิดลด) ** (-ปี)
    return ต้นทุน * pwf


def คำนวณต้นทุนเฉลี่ยรายปี(pw: float, อัตราคิดลด: float, ระยะวิเคราะห์: int) -> float:
    """
    แปลงมูลค่าปัจจุบันเป็นต้นทุนเฉลี่ยรายปี (EAC)
    สูตร: EAC = PW × [i(1+i)^N] / [(1+i)^N - 1]
    """
    if ระยะวิเคราะห์ <= 0 or อัตราคิดลด <= 0:
        return 0.0
    ตัวเศษ = อัตราคิดลด * (1 + อัตราคิดลด) ** ระยะวิเคราะห์
    ตัวส่วน = (1 + อัตราคิดลด) ** ระยะวิเคราะห์ - 1
    crf = ตัวเศษ / ตัวส่วน
    return pw * crf


def คำนวณมูลค่าซาก(
    ต้นทุนฟื้นฟูครั้งสุดท้าย: float,
    ปีฟื้นฟูครั้งสุดท้าย: int,
    อายุใช้งานที่คาดหวัง: int,
    ระยะวิเคราะห์: int,
    ร้อยละมูลค่าซาก: float = 20.0
) -> float:
    """คำนวณมูลค่าซากโดยวิธี Straight-Line Depreciation"""
    อายุใช้งานที่เหลือ = อายุใช้งานที่คาดหวัง - (ระยะวิเคราะห์ - ปีฟื้นฟูครั้งสุดท้าย)
    
    if อายุใช้งานที่เหลือ <= 0:
        return ต้นทุนฟื้นฟูครั้งสุดท้าย * (ร้อยละมูลค่าซาก / 100.0)
    
    ค่าเสื่อมต่อปี = ต้นทุนฟื้นฟูครั้งสุดท้าย * (1 - ร้อยละมูลค่าซาก/100.0) / อายุใช้งานที่คาดหวัง
    มูลค่าซาก = ต้นทุนฟื้นฟูครั้งสุดท้าย - ค่าเสื่อมต่อปี * (ระยะวิเคราะห์ - ปีฟื้นฟูครั้งสุดท้าย)
    
    return max(มูลค่าซาก, ต้นทุนฟื้นฟูครั้งสุดท้าย * ร้อยละมูลค่าซาก / 100.0)


# =============================================================================
# ส่วนที่ 3: สร้างตารางกระแสเงินสด
# =============================================================================

def สร้างตารางกระแสเงินสด(
    ทางเลือก: ทางเลือกผิวทาง,
    ระยะวิเคราะห์: int,
    อัตราคิดลด: float,
    รวมมูลค่าซาก: bool = True
) -> pd.DataFrame:
    """
    สร้างตารางกระแสเงินสดรายปี
    
    Logic แบบ C - รีเซ็ตรอบบำรุงรักษาหลังฟื้นฟูสภาพ:
    - เมื่อทำงานฟื้นฟู (Rehabilitation) ผิวทางเหมือนใหม่
    - รอบบำรุงรักษาเริ่มนับใหม่จากปีที่ทำฟื้นฟู
    - ไม่ทำบำรุงรักษาในปีเดียวกับฟื้นฟู
    """
    รายการ = []
    พื้นที่ = ทางเลือก.พื้นที่
    
    # เรียงลำดับปีฟื้นฟูสภาพ
    ปีฟื้นฟูทั้งหมด = sorted([ฟ.ปีดำเนินการ for ฟ in ทางเลือก.แผนฟื้นฟูสภาพ if ฟ.ปีดำเนินการ <= ระยะวิเคราะห์])
    ปีฟื้นฟู_set = set(ปีฟื้นฟูทั้งหมด)
    
    # ปีที่ 0: ต้นทุนก่อสร้างเริ่มต้น
    ต้นทุนเริ่มต้น = ทางเลือก.ต้นทุนก่อสร้าง * พื้นที่
    รายการ.append({
        'ปี': 0,
        'กิจกรรม': 'ก่อสร้างเริ่มต้น',
        'ประเภท': 'ก่อสร้าง',
        'ต้นทุนต่อหน่วย': ทางเลือก.ต้นทุนก่อสร้าง,
        'ต้นทุนตามปี': ต้นทุนเริ่มต้น,
        'ตัวคูณ_PW': 1.0,
        'มูลค่าปัจจุบัน': ต้นทุนเริ่มต้น
    })
    
    # กิจกรรมบำรุงรักษา (รีเซ็ตรอบหลังฟื้นฟู)
    for บำรุง in ทางเลือก.แผนบำรุงรักษา:
        if บำรุง.ความถี่ > 0:
            # สร้างช่วงเวลา: [0, ปีฟื้นฟู1, ปีฟื้นฟู2, ..., ระยะวิเคราะห์]
            จุดเริ่มต้นช่วง = [0] + ปีฟื้นฟูทั้งหมด
            
            for idx, ปีเริ่มช่วง in enumerate(จุดเริ่มต้นช่วง):
                # หาจุดสิ้นสุดช่วง
                if idx + 1 < len(จุดเริ่มต้นช่วง):
                    ปีสิ้นสุดช่วง = จุดเริ่มต้นช่วง[idx + 1]
                else:
                    ปีสิ้นสุดช่วง = ระยะวิเคราะห์ + 1
                
                # คำนวณปีบำรุงรักษาในช่วงนี้ (เริ่มนับจาก ปีเริ่มช่วง + ความถี่)
                ปี = ปีเริ่มช่วง + บำรุง.ความถี่
                while ปี < ปีสิ้นสุดช่วง and ปี <= ระยะวิเคราะห์:
                    # ข้ามถ้าตรงกับปีฟื้นฟู
                    if ปี not in ปีฟื้นฟู_set:
                        ต้นทุน = บำรุง.ต้นทุนต่อหน่วย * พื้นที่
                        pwf = (1 + อัตราคิดลด) ** (-ปี)
                        รายการ.append({
                            'ปี': ปี,
                            'กิจกรรม': บำรุง.ชื่อกิจกรรม,
                            'ประเภท': 'บำรุงรักษา',
                            'ต้นทุนต่อหน่วย': บำรุง.ต้นทุนต่อหน่วย,
                            'ต้นทุนตามปี': ต้นทุน,
                            'ตัวคูณ_PW': pwf,
                            'มูลค่าปัจจุบัน': ต้นทุน * pwf
                        })
                    ปี += บำรุง.ความถี่
    
    # กิจกรรมฟื้นฟูสภาพ
    for ฟื้นฟู in ทางเลือก.แผนฟื้นฟูสภาพ:
        if ฟื้นฟู.ปีดำเนินการ <= ระยะวิเคราะห์:
            ต้นทุน = ฟื้นฟู.ต้นทุนต่อหน่วย * พื้นที่
            pwf = (1 + อัตราคิดลด) ** (-ฟื้นฟู.ปีดำเนินการ)
            รายการ.append({
                'ปี': ฟื้นฟู.ปีดำเนินการ,
                'กิจกรรม': ฟื้นฟู.ชื่อกิจกรรม,
                'ประเภท': 'ฟื้นฟูสภาพ',
                'ต้นทุนต่อหน่วย': ฟื้นฟู.ต้นทุนต่อหน่วย,
                'ต้นทุนตามปี': ต้นทุน,
                'ตัวคูณ_PW': pwf,
                'มูลค่าปัจจุบัน': ต้นทุน * pwf
            })
    
    # มูลค่าซาก (ถ้าเปิดใช้งาน)
    if รวมมูลค่าซาก and len(ปีฟื้นฟูทั้งหมด) > 0:
        ปีฟื้นฟูสุดท้าย = max(ปีฟื้นฟูทั้งหมด)
        กิจกรรมฟื้นฟูสุดท้าย = [ฟ for ฟ in ทางเลือก.แผนฟื้นฟูสภาพ if ฟ.ปีดำเนินการ == ปีฟื้นฟูสุดท้าย]
        
        if len(กิจกรรมฟื้นฟูสุดท้าย) > 0:
            ต้นทุนฟื้นฟูสุดท้าย = sum([ฟ.ต้นทุนต่อหน่วย * พื้นที่ for ฟ in กิจกรรมฟื้นฟูสุดท้าย])
            อายุใช้งาน = 20  # สมมติอายุใช้งาน 20 ปี
            
            มูลค่าซาก = คำนวณมูลค่าซาก(
                ต้นทุนฟื้นฟูสุดท้าย,
                ปีฟื้นฟูสุดท้าย,
                อายุใช้งาน,
                ระยะวิเคราะห์,
                ทางเลือก.ร้อยละมูลค่าซาก
            )
            
            pwf = (1 + อัตราคิดลด) ** (-ระยะวิเคราะห์)
            รายการ.append({
                'ปี': ระยะวิเคราะห์,
                'กิจกรรม': 'มูลค่าซาก',
                'ประเภท': 'มูลค่าซาก',
                'ต้นทุนต่อหน่วย': -มูลค่าซาก / พื้นที่,
                'ต้นทุนตามปี': -มูลค่าซาก,
                'ตัวคูณ_PW': pwf,
                'มูลค่าปัจจุบัน': -มูลค่าซาก * pwf
            })
    
    # สร้าง DataFrame และเรียงตามปี
    df = pd.DataFrame(รายการ)
    df = df.sort_values('ปี').reset_index(drop=True)
    
    return df


# =============================================================================
# ส่วนที่ 4: ฟังก์ชันวิเคราะห์ LCCA
# =============================================================================

def วิเคราะห์_LCCA(
    ทางเลือกทั้งหมด: List[ทางเลือกผิวทาง],
    ระยะวิเคราะห์: int,
    อัตราคิดลด: float,
    รวมมูลค่าซาก: bool = True
) -> Tuple[pd.DataFrame, Dict[str, pd.DataFrame]]:
    """
    วิเคราะห์ LCCA สำหรับทางเลือกทั้งหมด
    
    Returns:
        สรุปผล (DataFrame), กระแสเงินสดแต่ละทางเลือก (Dict)
    """
    สรุป = []
    กระแสเงินสด = {}
    
    for ทางเลือก in ทางเลือกทั้งหมด:
        if not ทางเลือก.เปิดใช้งาน:
            continue
        
        # สร้างตารางกระแสเงินสด
        cf_table = สร้างตารางกระแสเงินสด(ทางเลือก, ระยะวิเคราะห์, อัตราคิดลด, รวมมูลค่าซาก)
        กระแสเงินสด[ทางเลือก.ชื่อ] = cf_table
        
        # คำนวณผลรวม
        มูลค่าปัจจุบันรวม = cf_table['มูลค่าปัจจุบัน'].sum()
        ต้นทุนเฉลี่ยรายปี = คำนวณต้นทุนเฉลี่ยรายปี(มูลค่าปัจจุบันรวม, อัตราคิดลด, ระยะวิเคราะห์)
        
        # คำนวณต้นทุนแยกตามประเภท
        ต้นทุนก่อสร้าง = cf_table[cf_table['ประเภท'] == 'ก่อสร้าง']['มูลค่าปัจจุบัน'].sum()
        ต้นทุนบำรุงรักษา = cf_table[cf_table['ประเภท'] == 'บำรุงรักษา']['มูลค่าปัจจุบัน'].sum()
        ต้นทุนฟื้นฟู = cf_table[cf_table['ประเภท'] == 'ฟื้นฟูสภาพ']['มูลค่าปัจจุบัน'].sum()
        มูลค่าซาก_pv = cf_table[cf_table['ประเภท'] == 'มูลค่าซาก']['มูลค่าปัจจุบัน'].sum()
        
        สรุป.append({
            'ทางเลือก': ทางเลือก.ชื่อ,
            'ประเภท': ทางเลือก.ประเภท,
            'ต้นทุนก่อสร้าง (PW)': ต้นทุนก่อสร้าง,
            'ต้นทุนบำรุงรักษา (PW)': ต้นทุนบำรุงรักษา,
            'ต้นทุนฟื้นฟู (PW)': ต้นทุนฟื้นฟู,
            'มูลค่าซาก (PW)': มูลค่าซาก_pv,
            'มูลค่าปัจจุบันรวม (PW)': มูลค่าปัจจุบันรวม,
            'ต้นทุนเฉลี่ยรายปี (EAC)': ต้นทุนเฉลี่ยรายปี
        })
    
    df_สรุป = pd.DataFrame(สรุป)
    
    return df_สรุป, กระแสเงินสด


def วิเคราะห์ความไว_อัตราคิดลด(
    ทางเลือกทั้งหมด: List[ทางเลือกผิวทาง],
    ระยะวิเคราะห์: int,
    อัตราคิดลดกลาง: float,
    ช่วงอัตราคิดลด: Tuple[float, float],
    รวมมูลค่าซาก: bool = True
) -> Tuple[pd.DataFrame, pd.DataFrame]:
    """วิเคราะห์ความไวต่ออัตราคิดลด"""
    อัตราต่ำสุด, อัตราสูงสุด = ช่วงอัตราคิดลด
    รายการอัตรา = np.arange(อัตราต่ำสุด, อัตราสูงสุด + 0.005, 0.01)
    
    ผลลัพธ์ = []
    
    for อัตรา in รายการอัตรา:
        สรุป, _ = วิเคราะห์_LCCA(ทางเลือกทั้งหมด, ระยะวิเคราะห์, อัตรา, รวมมูลค่าซาก)
        
        for _, row in สรุป.iterrows():
            ผลลัพธ์.append({
                'อัตราคิดลด': อัตรา,
                'ทางเลือก': row['ทางเลือก'],
                'มูลค่าปัจจุบัน': row['มูลค่าปัจจุบันรวม (PW)']
            })
    
    df_ผล = pd.DataFrame(ผลลัพธ์)
    
    # สร้าง pivot table
    pivot = df_ผล.pivot(index='อัตราคิดลด', columns='ทางเลือก', values='มูลค่าปัจจุบัน')
    
    return df_ผล, pivot


# =============================================================================
# ส่วนที่ 5: Excel Template และ Upload Functions
# =============================================================================

def สร้าง_excel_template() -> io.BytesIO:
    """สร้าง Excel template สำหรับผู้ใช้กรอกข้อมูล"""
    
    # สร้างข้อมูลตัวอย่าง
    data = {
        'ผิวทาง': ['AC', 'JPCP', 'JRCP', 'CRCP'],
        'ประเภทผิวทาง': ['ลาดยาง', 'คอนกรีต', 'คอนกรีต', 'คอนกรีต'],
        'ความหนาผิวทาง (ซม.)': ['กรอกข้อมูล', 'กรอกข้อมูล', 'กรอกข้อมูล', 'กรอกข้อมูล'],
        'ต้นทุนก่อสร้าง (บาท/ตร.ม.)': ['กรอกข้อมูล', 'กรอกข้อมูล', 'กรอกข้อมูล', 'กรอกข้อมูล']
    }
    
    df = pd.DataFrame(data)
    
    # สร้าง Excel file
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        # เขียน header
        header_df = pd.DataFrame([['ข้อมูลสำหรับวิเคราะห์ LCCA']])
        header_df.to_excel(writer, sheet_name='Sheet1', index=False, header=False)
        
        # เขียนข้อมูล
        df.to_excel(writer, sheet_name='Sheet1', index=False, startrow=1)
        
        # จัดรูปแบบ
        worksheet = writer.sheets['Sheet1']
        
        # ตั้งค่าความกว้างคอลัมน์
        worksheet.column_dimensions['A'].width = 15
        worksheet.column_dimensions['B'].width = 20
        worksheet.column_dimensions['C'].width = 25
        worksheet.column_dimensions['D'].width = 30
    
    output.seek(0)
    return output


def อ่านข้อมูลจาก_excel(uploaded_file) -> Dict[str, float]:
    """
    อ่านข้อมูลจากไฟล์ Excel ที่ผู้ใช้ upload
    
    Returns:
        Dict mapping ชื่อผิวทาง -> ต้นทุนก่อสร้าง
    """
    try:
        # อ่านไฟล์ Excel โดยข้าม header row
        df = pd.read_excel(uploaded_file, sheet_name='Sheet1', header=1)
        
        # สร้าง dictionary
        ข้อมูลต้นทุน = {}
        
        for idx, row in df.iterrows():
            ชื่อ = row['ผิวทาง']
            ต้นทุน_str = str(row['ต้นทุนก่อสร้าง (บาท/ตร.ม.)'])
            
            # แปลงเป็นตัวเลข
            try:
                ต้นทุน = float(ต้นทุน_str.replace(',', ''))
                ข้อมูลต้นทุน[ชื่อ] = ต้นทุน
            except:
                continue
        
        return ข้อมูลต้นทุน
    
    except Exception as e:
        st.error(f"❌ เกิดข้อผิดพลาดในการอ่านไฟล์: {str(e)}")
        return {}


# =============================================================================
# ส่วนที่ 6: ฟังก์ชันสร้างรายงาน Word
# =============================================================================

def สร้างรายงาน_word(
    ชื่อโครงการ: str,
    สรุป: pd.DataFrame,
    กระแสเงินสด: Dict[str, pd.DataFrame],
    ระยะวิเคราะห์: int,
    อัตราคิดลด: float,
    รวมมูลค่าซาก: bool
) -> io.BytesIO:
    """สร้างรายงาน Word"""
    
    if not DOCX_AVAILABLE:
        return None
    
    doc = WordDocument()
    
    # ตั้งค่า font ภาษาไทย
    style = doc.styles['Normal']
    font = style.font
    font.name = 'TH Sarabun New'
    font.size = Pt(16)
    
    # หัวเรื่อง
    title = doc.add_heading('รายงานการวิเคราะห์ต้นทุนตลอดอายุการใช้งาน', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ข้อมูลโครงการ
    doc.add_heading('ข้อมูลโครงการ', 1)
    info = doc.add_paragraph()
    info.add_run(f'ชื่อโครงการ: ').bold = True
    info.add_run(f'{ชื่อโครงการ}\n')
    info.add_run(f'ระยะเวลาวิเคราะห์: ').bold = True
    info.add_run(f'{ระยะวิเคราะห์} ปี\n')
    info.add_run(f'อัตราคิดลด: ').bold = True
    info.add_run(f'{อัตราคิดลด*100:.1f}%\n')
    info.add_run(f'รวมมูลค่าซาก: ').bold = True
    info.add_run(f'{"ใช่" if รวมมูลค่าซาก else "ไม่"}\n')
    info.add_run(f'วันที่จัดทำ: ').bold = True
    info.add_run(f'{datetime.now().strftime("%d/%m/%Y %H:%M")}\n')
    
    # ตารางสรุป
    doc.add_heading('สรุปผลการวิเคราะห์', 1)
    
    # สร้างตาราง
    table = doc.add_table(rows=1, cols=len(สรุป.columns))
    table.style = 'Light Grid Accent 1'
    
    # Header
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(สรุป.columns):
        hdr_cells[i].text = col_name
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Data
    for _, row in สรุป.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            if isinstance(value, (int, float)) and i > 1:  # ตัวเลข
                row_cells[i].text = f'{value:,.0f}'
            else:
                row_cells[i].text = str(value)
    
    # กระแสเงินสดแต่ละทางเลือก
    doc.add_page_break()
    doc.add_heading('รายละเอียดกระแสเงินสด', 1)
    
    for ชื่อทางเลือก, cf_df in กระแสเงินสด.items():
        doc.add_heading(f'ทางเลือก: {ชื่อทางเลือก}', 2)
        
        # สร้างตาราง
        cf_table = doc.add_table(rows=1, cols=len(cf_df.columns))
        cf_table.style = 'Light List Accent 1'
        
        # Header
        hdr_cells = cf_table.rows[0].cells
        for i, col_name in enumerate(cf_df.columns):
            hdr_cells[i].text = col_name
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Data (แสดงแค่ 20 แถวแรก)
        for idx, (_, row) in enumerate(cf_df.head(20).iterrows()):
            row_cells = cf_table.add_row().cells
            for i, value in enumerate(row):
                if isinstance(value, float):
                    if i == 5:  # ตัวคูณ PW
                        row_cells[i].text = f'{value:.4f}'
                    else:
                        row_cells[i].text = f'{value:,.2f}'
                else:
                    row_cells[i].text = str(value)
        
        if len(cf_df) > 20:
            doc.add_paragraph(f'... (แสดง 20 จาก {len(cf_df)} แถว)')
        
        doc.add_paragraph()
    
    # บันทึกลง BytesIO
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output


# =============================================================================
# ส่วนที่ 7: ฟังก์ชันหลัก (Main Function)
# =============================================================================

def main():
    """ฟังก์ชันหลักของโปรแกรม"""
    
    # Custom CSS
    st.markdown("""
    <style>
    .main-header {
        font-size: 2.5rem;
        color: #1f77b4;
        text-align: center;
        font-weight: bold;
        margin-bottom: 1rem;
    }
    .metric-card {
        background-color: #f0f2f6;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 0.5rem 0;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # หัวเรื่อง
    st.markdown('<div class="main-header">🛣️ โปรแกรมวิเคราะห์ LCCA ผิวทาง v2.1</div>', unsafe_allow_html=True)
    st.markdown("**Life-Cycle Cost Analysis for Pavement Design**")
    
    # ==========================================================================
    # Sidebar: การตั้งค่าและ Upload
    # ==========================================================================
    
    with st.sidebar:
        st.header("⚙️ การตั้งค่าโครงการ")
        
        # ข้อมูลพื้นฐาน
        if 'ชื่อโครงการ' not in st.session_state:
            st.session_state.ชื่อโครงการ = "โครงการทางหลวงสายหลัก"
        
        ชื่อโครงการ = st.text_input(
            "ชื่อโครงการ:",
            value=st.session_state.ชื่อโครงการ,
            key="input_project_name"
        )
        st.session_state.ชื่อโครงการ = ชื่อโครงการ
        
        พื้นที่โครงการ = st.number_input(
            "พื้นที่โครงการ (ตร.ม.):",
            min_value=100.0,
            max_value=1000000.0,
            value=10000.0,
            step=1000.0
        )
        
        st.divider()
        
        # พารามิเตอร์การวิเคราะห์
        st.subheader("📊 พารามิเตอร์การวิเคราะห์")
        
        ระยะวิเคราะห์ = st.slider(
            "ระยะเวลาวิเคราะห์ (ปี):",
            min_value=10,
            max_value=50,
            value=30,
            step=5
        )
        
        อัตราคิดลด = st.slider(
            "อัตราคิดลด (%):",
            min_value=0.0,
            max_value=10.0,
            value=3.0,
            step=0.5
        ) / 100.0
        
        st.divider()
        
        # Toggle มูลค่าซาก
        st.subheader("💰 มูลค่าซาก (Salvage Value)")
        
        รวมมูลค่าซาก = st.toggle(
            "รวมมูลค่าซากในการคำนวณ",
            value=True,
            help="เปิดใช้งานเพื่อนำมูลค่าซากมาหักออกจากต้นทุนทั้งหมด"
        )
        
        if รวมมูลค่าซาก:
            st.info("✅ กำลังคำนวณมูลค่าซาก")
        else:
            st.warning("⚠️ ไม่คำนวณมูลค่าซาก")
        
        st.divider()
        
        # Excel Upload Section
        st.subheader("📤 อัปโหลดข้อมูลต้นทุน")
        
        # ดาวน์โหลด Template
        template_file = สร้าง_excel_template()
        st.download_button(
            label="📥 ดาวน์โหลด Template Excel",
            data=template_file,
            file_name="LCCA_Template.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.document",
            use_container_width=True
        )
        
        st.caption("💡 ดาวน์โหลด template → กรอกข้อมูล → อัปโหลดกลับมา")
        
        # Upload File
        uploaded_file = st.file_uploader(
            "เลือกไฟล์ Excel:",
            type=['xlsx', 'xls'],
            help="อัปโหลดไฟล์ Excel ที่กรอกข้อมูลแล้ว"
        )
        
        # Process uploaded file
        if uploaded_file is not None:
            try:
                # อ่านข้อมูล
                ข้อมูลต้นทุน = อ่านข้อมูลจาก_excel(uploaded_file)
                
                if len(ข้อมูลต้นทุน) > 0:
                    # สร้าง version hash
                    file_hash = hashlib.md5(uploaded_file.getvalue()).hexdigest()[:8]
                    
                    # บันทึกลง session_state
                    st.session_state['uploaded_cost_data'] = ข้อมูลต้นทุน
                    st.session_state['upload_version'] = file_hash
                    
                    st.success(f"✅ อ่านข้อมูลสำเร็จ! ({len(ข้อมูลต้นทุน)} รายการ)")
                    
                    # แสดงข้อมูลที่อ่านได้
                    with st.expander("👀 ดูข้อมูลที่อัปโหลด"):
                        for ชื่อ, ต้นทุน in ข้อมูลต้นทุน.items():
                            st.write(f"• {ชื่อ}: {ต้นทุน:,.0f} บาท/ตร.ม.")
                else:
                    st.error("❌ ไม่พบข้อมูลที่ถูกต้องในไฟล์")
            
            except Exception as e:
                st.error(f"❌ เกิดข้อผิดพลาด: {str(e)}")
        
        st.divider()
        
        # ช่วงการวิเคราะห์ความไว
        st.subheader("📈 การวิเคราะห์ความไว")
        col1, col2 = st.columns(2)
        with col1:
            อัตราต่ำสุด = st.number_input("อัตราต่ำสุด (%):", value=1.0, step=0.5) / 100.0
        with col2:
            อัตราสูงสุด = st.number_input("อัตราสูงสุด (%):", value=8.0, step=0.5) / 100.0
        
        ช่วงอัตราคิดลด = (อัตราต่ำสุด, อัตราสูงสุด)
    
    # ==========================================================================
    # Main Content: จัดการทางเลือก
    # ==========================================================================
    
    # Initialize session state
    if 'ทางเลือกทั้งหมด' not in st.session_state:
        # ค่าเริ่มต้น
        st.session_state.ทางเลือกทั้งหมด = [
            ทางเลือกผิวทาง(
                ชื่อ="AC",
                ประเภท="Flexible (แอสฟัลต์)",
                ต้นทุนก่อสร้าง=800.0,
                แผนบำรุงรักษา=[
                    กิจกรรมบำรุงรักษา("Crack Sealing", 50.0, 0, 3),
                    กิจกรรมบำรุงรักษา("Surface Treatment", 150.0, 0, 5)
                ],
                แผนฟื้นฟูสภาพ=[
                    กิจกรรมฟื้นฟูสภาพ("Overlay (5 cm)", 400.0, 12),
                    กิจกรรมฟื้นฟูสภาพ("Overlay (5 cm)", 400.0, 24)
                ],
                พื้นที่=พื้นที่โครงการ,
                ความหนา=10.0
            ),
            ทางเลือกผิวทาง(
                ชื่อ="JPCP",
                ประเภท="Rigid (คอนกรีตธรรมดา)",
                ต้นทุนก่อสร้าง=1200.0,
                แผนบำรุงรักษา=[
                    กิจกรรมบำรุงรักษา("Joint Sealing", 30.0, 0, 5),
                    กิจกรรมบำรุงรักษา("Slab Repair", 100.0, 0, 10)
                ],
                แผนฟื้นฟูสภาพ=[
                    กิจกรรมฟื้นฟูสภาพ("Diamond Grinding", 150.0, 15)
                ],
                พื้นที่=พื้นที่โครงการ,
                ความหนา=25.0
            ),
            ทางเลือกผิวทาง(
                ชื่อ="JRCP",
                ประเภท="Rigid (คอนกรีตเสริมเหล็ก)",
                ต้นทุนก่อสร้าง=1400.0,
                แผนบำรุงรักษา=[
                    กิจกรรมบำรุงรักษา("Joint Sealing", 30.0, 0, 5),
                    กิจกรรมบำรุงรักษา("Partial Repair", 80.0, 0, 12)
                ],
                แผนฟื้นฟูสภาพ=[
                    กิจกรรมฟื้นฟูสภาพ("Diamond Grinding", 150.0, 18)
                ],
                พื้นที่=พื้นที่โครงการ,
                ความหนา=28.0
            ),
            ทางเลือกผิวทาง(
                ชื่อ="CRCP",
                ประเภท="Rigid (คอนกรีตต่อเนื่อง)",
                ต้นทุนก่อสร้าง=1600.0,
                แผนบำรุงรักษา=[
                    กิจกรรมบำรุงรักษา("Surface Cleaning", 20.0, 0, 5),
                    กิจกรรมบำรุงรักษา("Minor Repair", 60.0, 0, 15)
                ],
                แผนฟื้นฟูสภาพ=[
                    กิจกรรมฟื้นฟูสภาพ("Diamond Grinding", 150.0, 20)
                ],
                พื้นที่=พื้นที่โครงการ,
                ความหนา=30.0
            )
        ]
    
    # อัพเดทพื้นที่ทุกทางเลือก
    for ทางเลือก in st.session_state.ทางเลือกทั้งหมด:
        ทางเลือก.พื้นที่ = พื้นที่โครงการ
    
    # Get upload version for dynamic keys
    upload_version = st.session_state.get('upload_version', 'default')
    uploaded_cost = st.session_state.get('uploaded_cost_data', {})
    
    # ==========================================================================
    # Tabs
    # ==========================================================================
    
    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "🏗️ จัดการทางเลือก",
        "📊 ผลการวิเคราะห์",
        "💰 กระแสเงินสด",
        "📈 การวิเคราะห์ความไว",
        "ℹ️ ทฤษฎี LCCA"
    ])
    
    # ==========================================================================
    # แท็บ 1: จัดการทางเลือก
    # ==========================================================================
    with tab1:
        st.header("🏗️ จัดการทางเลือกผิวทาง")
        
        # เลือกทางเลือกที่จะแก้ไข
        ชื่อทางเลือกทั้งหมด = [ท.ชื่อ for ท in st.session_state.ทางเลือกทั้งหมด]
        
        ทางเลือกที่เลือก_idx = st.selectbox(
            "เลือกทางเลือกที่ต้องการแก้ไข:",
            options=range(len(st.session_state.ทางเลือกทั้งหมด)),
            format_func=lambda x: st.session_state.ทางเลือกทั้งหมด[x].ชื่อ
        )
        
        ทางเลือก = st.session_state.ทางเลือกทั้งหมด[ทางเลือกที่เลือก_idx]
        
        st.divider()
        
        # แก้ไขข้อมูลพื้นฐาน
        col1, col2 = st.columns(2)
        
        with col1:
            # ถ้ามีข้อมูล upload ให้ใช้ข้อมูลนั้น
            ต้นทุนเริ่มต้น = uploaded_cost.get(ทางเลือก.ชื่อ, ทางเลือก.ต้นทุนก่อสร้าง)
            
            ทางเลือก.ต้นทุนก่อสร้าง = st.number_input(
                "ต้นทุนก่อสร้าง (บาท/ตร.ม.):",
                min_value=0.0,
                max_value=10000.0,
                value=float(ต้นทุนเริ่มต้น),
                step=50.0,
                key=f"cost_{ทางเลือก.ชื่อ}_{upload_version}"  # Dynamic key
            )
            
            ทางเลือก.ความหนา = st.number_input(
                "ความหนาผิวทาง (ซม.):",
                min_value=0.0,
                max_value=50.0,
                value=float(ทางเลือก.ความหนา),
                step=1.0,
                key=f"thick_{ทางเลือก.ชื่อ}_{upload_version}"
            )
        
        with col2:
            ทางเลือก.ร้อยละมูลค่าซาก = st.number_input(
                "ร้อยละมูลค่าซาก (%):",
                min_value=0.0,
                max_value=50.0,
                value=float(ทางเลือก.ร้อยละมูลค่าซาก),
                step=5.0,
                key=f"salvage_{ทางเลือก.ชื่อ}_{upload_version}"
            )
            
            ทางเลือก.เปิดใช้งาน = st.checkbox(
                "เปิดใช้งานทางเลือกนี้",
                value=ทางเลือก.เปิดใช้งาน,
                key=f"enable_{ทางเลือก.ชื่อ}_{upload_version}"
            )
        
        st.divider()
        
        # แสดงข้อมูลสรุป
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("ต้นทุนก่อสร้างรวม", f"{ทางเลือก.ต้นทุนก่อสร้าง * พื้นที่โครงการ:,.0f} บาท")
        with col2:
            st.metric("จำนวนกิจกรรมบำรุงรักษา", len(ทางเลือก.แผนบำรุงรักษา))
        with col3:
            st.metric("จำนวนกิจกรรมฟื้นฟู", len(ทางเลือก.แผนฟื้นฟูสภาพ))
        
        st.divider()
        
        # แสดงแผนบำรุงรักษาและฟื้นฟู
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("🔧 แผนบำรุงรักษา")
            if len(ทางเลือก.แผนบำรุงรักษา) > 0:
                for บำรุง in ทางเลือก.แผนบำรุงรักษา:
                    with st.expander(f"📌 {บำรุง.ชื่อกิจกรรม}"):
                        st.write(f"**ต้นทุน:** {บำรุง.ต้นทุนต่อหน่วย:,.0f} บาท/ตร.ม.")
                        st.write(f"**ความถี่:** ทุก {บำรุง.ความถี่} ปี")
            else:
                st.info("ไม่มีกิจกรรมบำรุงรักษา")
        
        with col2:
            st.subheader("🔄 แผนฟื้นฟูสภาพ")
            if len(ทางเลือก.แผนฟื้นฟูสภาพ) > 0:
                for ฟื้นฟู in ทางเลือก.แผนฟื้นฟูสภาพ:
                    with st.expander(f"📌 {ฟื้นฟู.ชื่อกิจกรรม}"):
                        st.write(f"**ต้นทุน:** {ฟื้นฟู.ต้นทุนต่อหน่วย:,.0f} บาท/ตร.ม.")
                        st.write(f"**ปีที่ดำเนินการ:** ปีที่ {ฟื้นฟู.ปีดำเนินการ}")
            else:
                st.info("ไม่มีกิจกรรมฟื้นฟูสภาพ")
    
    # ==========================================================================
    # แท็บ 2: ผลการวิเคราะห์
    # ==========================================================================
    with tab2:
        st.header("📊 ผลการวิเคราะห์ LCCA")
        
        ทางเลือกที่ใช้ = [ท for ท in st.session_state.ทางเลือกทั้งหมด if ท.เปิดใช้งาน]
        
        if len(ทางเลือกที่ใช้) == 0:
            st.warning("⚠️ กรุณาเปิดใช้งานอย่างน้อย 1 ทางเลือก")
        else:
            # วิเคราะห์
            สรุป, กระแสเงินสด = วิเคราะห์_LCCA(
                st.session_state.ทางเลือกทั้งหมด,
                ระยะวิเคราะห์,
                อัตราคิดลด,
                รวมมูลค่าซาก
            )
            
            # หาทางเลือกที่ดีที่สุด
            ทางเลือกที่ดีที่สุด = สรุป.loc[สรุป['มูลค่าปัจจุบันรวม (PW)'].idxmin()]
            
            # แสดงผลแบบเด่น
            st.success(f"""
            ### 🏆 ทางเลือกที่ดีที่สุด: **{ทางเลือกที่ดีที่สุด['ทางเลือก']}**
            - มูลค่าปัจจุบันรวม: **{ทางเลือกที่ดีที่สุด['มูลค่าปัจจุบันรวม (PW)']:,.0f}** บาท
            - ต้นทุนเฉลี่ยรายปี: **{ทางเลือกที่ดีที่สุด['ต้นทุนเฉลี่ยรายปี (EAC)']:,.0f}** บาท/ปี
            """)
            
            st.divider()
            
            # Metrics
            col1, col2, col3, col4 = st.columns(4)
            
            for idx, (_, row) in enumerate(สรุป.iterrows()):
                with [col1, col2, col3, col4][idx % 4]:
                    delta = None
                    if idx > 0:
                        ผลต่างจากที่ดีที่สุด = row['มูลค่าปัจจุบันรวม (PW)'] - ทางเลือกที่ดีที่สุด['มูลค่าปัจจุบันรวม (PW)']
                        delta = f"+{ผลต่างจากที่ดีที่สุด:,.0f}"
                    
                    st.metric(
                        label=f"{row['ทางเลือก']}",
                        value=f"{row['มูลค่าปัจจุบันรวม (PW)']:,.0f} บาท",
                        delta=delta,
                        delta_color="inverse"
                    )
            
            st.divider()
            
            # ตารางสรุป
            st.subheader("📋 ตารางสรุปผลการวิเคราะห์")
            
            สรุป_display = สรุป.copy()
            for col in สรุป_display.columns[2:]:
                สรุป_display[col] = สรุป_display[col].apply(lambda x: f"{x:,.0f}")
            
            st.dataframe(สรุป_display, use_container_width=True, hide_index=True)
            
            st.divider()
            
            # กราฟเปรียบเทียบ
            st.subheader("📊 กราฟเปรียบเทียบต้นทุน")
            
            # กราฟแท่ง - แยกตามประเภทต้นทุน
            fig_breakdown = go.Figure()
            
            categories = ['ก่อสร้าง', 'บำรุงรักษา', 'ฟื้นฟู', 'มูลค่าซาก']
            colors = ['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728']
            
            for idx, cat in enumerate(categories):
                col_name = f'ต้นทุน{cat} (PW)' if cat != 'มูลค่าซาก' else 'มูลค่าซาก (PW)'
                fig_breakdown.add_trace(go.Bar(
                    name=cat,
                    x=สรุป['ทางเลือก'],
                    y=สรุป[col_name],
                    marker_color=colors[idx]
                ))
            
            fig_breakdown.update_layout(
                title='การแยกประเภทต้นทุนตามมูลค่าปัจจุบัน',
                xaxis_title='ทางเลือก',
                yaxis_title='มูลค่าปัจจุบัน (บาท)',
                barmode='relative',
                height=500
            )
            
            st.plotly_chart(fig_breakdown, use_container_width=True)
            
            # กราฟวงกลม - สัดส่วนต้นทุน
            col1, col2 = st.columns(2)
            
            with col1:
                fig_pie_pw = px.pie(
                    สรุป,
                    values='มูลค่าปัจจุบันรวม (PW)',
                    names='ทางเลือก',
                    title='สัดส่วนมูลค่าปัจจุบัน'
                )
                st.plotly_chart(fig_pie_pw, use_container_width=True)
            
            with col2:
                fig_pie_eac = px.pie(
                    สรุป,
                    values='ต้นทุนเฉลี่ยรายปี (EAC)',
                    names='ทางเลือก',
                    title='สัดส่วนต้นทุนเฉลี่ยรายปี'
                )
                st.plotly_chart(fig_pie_eac, use_container_width=True)
            
            st.divider()
            
            # ส่งออกรายงาน
            st.subheader("📥 ส่งออกรายงาน")
            
            col_export1, col_export2 = st.columns(2)
            
            with col_export1:
                if DOCX_AVAILABLE:
                    word_file = สร้างรายงาน_word(
                        ชื่อโครงการ,
                        สรุป,
                        กระแสเงินสด,
                        ระยะวิเคราะห์,
                        อัตราคิดลด,
                        รวมมูลค่าซาก
                    )
                    st.download_button(
                        label="📝 ดาวน์โหลดรายงาน Word",
                        data=word_file,
                        file_name=f"LCCA_{st.session_state.ชื่อโครงการ}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                else:
                    st.warning("⚠️ ต้องติดตั้ง python-docx: `pip install python-docx`")
            
            with col_export2:
                csv_summary = สรุป.to_csv(index=False).encode('utf-8-sig')
                st.download_button(
                    label="📊 ดาวน์โหลดสรุป CSV",
                    data=csv_summary,
                    file_name=f"LCCA_Summary_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                    mime="text/csv",
                    use_container_width=True
                )
    
    # ==========================================================================
    # แท็บ 3: กระแสเงินสด
    # ==========================================================================
    with tab3:
        st.header("💰 ตารางกระแสเงินสดรายปี")
        
        ทางเลือกที่ใช้ = [ท for ท in st.session_state.ทางเลือกทั้งหมด if ท.เปิดใช้งาน]
        
        if len(ทางเลือกที่ใช้) == 0:
            st.warning("⚠️ กรุณาเปิดใช้งานอย่างน้อย 1 ทางเลือก")
        else:
            สรุป, กระแสเงินสด = วิเคราะห์_LCCA(
                st.session_state.ทางเลือกทั้งหมด,
                ระยะวิเคราะห์,
                อัตราคิดลด,
                รวมมูลค่าซาก
            )
            
            ทางเลือกที่เลือก = st.selectbox(
                "เลือกทางเลือกที่ต้องการดูรายละเอียด:",
                options=[ท.ชื่อ for ท in ทางเลือกที่ใช้]
            )
            
            if ทางเลือกที่เลือก in กระแสเงินสด:
                cf_table = กระแสเงินสด[ทางเลือกที่เลือก].copy()
                
                # แสดงสรุป
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("ต้นทุนตามปีรวม", f"{cf_table['ต้นทุนตามปี'].sum():,.0f} บาท")
                with col2:
                    st.metric("มูลค่าปัจจุบันรวม", f"{cf_table['มูลค่าปัจจุบัน'].sum():,.0f} บาท")
                with col3:
                    eac = คำนวณต้นทุนเฉลี่ยรายปี(cf_table['มูลค่าปัจจุบัน'].sum(), อัตราคิดลด, ระยะวิเคราะห์)
                    st.metric("EAC", f"{eac:,.0f} บาท/ปี")
                
                st.divider()
                
                # จัดรูปแบบตาราง
                cf_display = cf_table.copy()
                cf_display['ต้นทุนต่อหน่วย'] = cf_display['ต้นทุนต่อหน่วย'].apply(lambda x: f"{x:,.2f}")
                cf_display['ต้นทุนตามปี'] = cf_display['ต้นทุนตามปี'].apply(lambda x: f"{x:,.0f}")
                cf_display['ตัวคูณ_PW'] = cf_display['ตัวคูณ_PW'].apply(lambda x: f"{x:.4f}")
                cf_display['มูลค่าปัจจุบัน'] = cf_display['มูลค่าปัจจุบัน'].apply(lambda x: f"{x:,.0f}")
                cf_display.columns = ['ปี', 'กิจกรรม', 'ประเภท', 'ต้นทุน/หน่วย', 'ต้นทุนตามปี (บาท)', 'ตัวคูณ PW', 'มูลค่าปัจจุบัน (บาท)']
                
                st.dataframe(cf_display, use_container_width=True, hide_index=True, height=500)
                
                # กราฟ Timeline
                st.subheader("📅 Timeline กระแสเงินสด")
                
                cf_plot = cf_table[cf_table['ต้นทุนตามปี'] > 0].copy()
                
                fig_timeline = px.scatter(
                    cf_plot,
                    x='ปี',
                    y='มูลค่าปัจจุบัน',
                    size='ต้นทุนตามปี',
                    color='ประเภท',
                    hover_name='กิจกรรม',
                    title=f'Timeline กระแสเงินสด - {ทางเลือกที่เลือก}',
                    labels={'ปี': 'ปี', 'มูลค่าปัจจุบัน': 'มูลค่าปัจจุบัน (บาท)'}
                )
                fig_timeline.update_layout(height=400)
                st.plotly_chart(fig_timeline, use_container_width=True)
                
                # ดาวน์โหลด CSV
                st.divider()
                csv = cf_table.to_csv(index=False).encode('utf-8-sig')
                st.download_button(
                    label="⬇️ ดาวน์โหลดตารางกระแสเงินสด (CSV)",
                    data=csv,
                    file_name=f"cashflow_{ทางเลือกที่เลือก}.csv",
                    mime="text/csv"
                )
    
    # ==========================================================================
    # แท็บ 4: การวิเคราะห์ความไว
    # ==========================================================================
    with tab4:
        st.header("📈 การวิเคราะห์ความไว (Sensitivity Analysis)")
        
        ทางเลือกที่ใช้ = [ท for ท in st.session_state.ทางเลือกทั้งหมด if ท.เปิดใช้งาน]
        
        if len(ทางเลือกที่ใช้) == 0:
            st.warning("⚠️ กรุณาเปิดใช้งานอย่างน้อย 1 ทางเลือก")
        else:
            st.subheader("1️⃣ ความไวต่ออัตราคิดลด")
            
            ผลอัตราคิดลด, pivot_อัตราคิดลด = วิเคราะห์ความไว_อัตราคิดลด(
                st.session_state.ทางเลือกทั้งหมด,
                ระยะวิเคราะห์,
                อัตราคิดลด,
                ช่วงอัตราคิดลด,
                รวมมูลค่าซาก
            )
            
            if len(ผลอัตราคิดลด) > 0:
                # กราฟเส้น
                fig_sens = px.line(
                    ผลอัตราคิดลด,
                    x='อัตราคิดลด',
                    y='มูลค่าปัจจุบัน',
                    color='ทางเลือก',
                    markers=True,
                    title='ผลกระทบของอัตราคิดลดต่อมูลค่าปัจจุบัน',
                    labels={'อัตราคิดลด': 'อัตราคิดลด', 'มูลค่าปัจจุบัน': 'มูลค่าปัจจุบัน (บาท)'}
                )
                fig_sens.update_layout(height=500)
                fig_sens.update_xaxes(tickformat='.1%')
                st.plotly_chart(fig_sens, use_container_width=True)
                
                # ตาราง Pivot
                st.markdown("**ตารางสรุปมูลค่าปัจจุบันตามอัตราคิดลด (บาท):**")
                pivot_display = pivot_อัตราคิดลด.copy()
                for col in pivot_display.columns:
                    pivot_display[col] = pivot_display[col].apply(lambda x: f"{x:,.0f}")
                st.dataframe(pivot_display, use_container_width=True)
                
                # วิเคราะห์
                อัตราต่ำสุด = ผลอัตราคิดลด['อัตราคิดลด'].min()
                อัตราสูงสุด = ผลอัตราคิดลด['อัตราคิดลด'].max()
                
                ผู้ชนะต่ำ = ผลอัตราคิดลด[ผลอัตราคิดลด['อัตราคิดลด'] == อัตราต่ำสุด].nsmallest(1, 'มูลค่าปัจจุบัน')['ทางเลือก'].values[0]
                ผู้ชนะสูง = ผลอัตราคิดลด[ผลอัตราคิดลด['อัตราคิดลด'] == อัตราสูงสุด].nsmallest(1, 'มูลค่าปัจจุบัน')['ทางเลือก'].values[0]
                
                if ผู้ชนะต่ำ == ผู้ชนะสูง:
                    st.success(f"✅ **{ผู้ชนะต่ำ}** เป็นทางเลือกที่ประหยัดที่สุดในทุกอัตราคิดลด (Robust Decision)")
                else:
                    st.warning(f"⚠️ ทางเลือกที่ดีที่สุดเปลี่ยนแปลง: {ผู้ชนะต่ำ} (อัตราต่ำ) vs {ผู้ชนะสูง} (อัตราสูง)")
    
    # ==========================================================================
    # แท็บ 5: ทฤษฎี LCCA
    # ==========================================================================
    with tab5:
        st.header("ℹ️ ทฤษฎี Life-Cycle Cost Analysis (LCCA)")
        
        st.markdown("""
        ## 1. ประเภทผิวทางคอนกรีต
        
        | ประเภท | ชื่อเต็ม | ลักษณะเด่น |
        |--------|---------|-----------|
        | **JPCP** | Jointed Plain Concrete Pavement | คอนกรีตไม่เสริมเหล็ก มีรอยต่อทุก 4-6 ม. |
        | **JRCP** | Jointed Reinforced Concrete Pavement | คอนกรีตเสริมเหล็ก รอยต่อห่าง 8-15 ม. |
        | **CRCP** | Continuously Reinforced Concrete Pavement | เสริมเหล็กต่อเนื่อง ไม่มีรอยต่อตามขวาง |
        
        ## 2. สูตรคำนวณหลัก
        
        ### 2.1 มูลค่าปัจจุบัน (Present Worth)
        """)
        
        st.latex(r"PW = FV \times (1 + i)^{-n}")
        
        st.markdown("""
        ### 2.2 ต้นทุนเฉลี่ยรายปี (Equivalent Annual Cost)
        """)
        
        st.latex(r"EAC = PW \times \frac{i(1+i)^N}{(1+i)^N - 1}")
        
        st.markdown("""
        ## 3. เปรียบเทียบผิวทางคอนกรีต
        
        | เกณฑ์ | JPCP | JRCP | CRCP |
        |------|------|------|------|
        | ต้นทุนก่อสร้าง | ต่ำ | ปานกลาง | สูง |
        | ระยะห่างรอยต่อ | 4-6 ม. | 8-15 ม. | ไม่มี |
        | เหล็กเสริม | ไม่มี | 0.1-0.25% | 0.6-0.7% |
        | ค่าบำรุงรักษา | สูง | ปานกลาง | ต่ำ |
        | อายุใช้งาน | 20-30 ปี | 25-35 ปี | 30-40 ปี |
        
        ## 4. มูลค่าซาก (Salvage Value)
        
        มูลค่าซากคือมูลค่าที่เหลืออยู่ของผิวทางเมื่อสิ้นสุดระยะเวลาวิเคราะห์ 
        การคำนวณมูลค่าซากช่วยให้การเปรียบเทียบยุติธรรมมากขึ้น โดยเฉพาะเมื่อ:
        
        - ผิวทางแต่ละประเภทมีอายุการใช้งานที่แตกต่างกัน
        - มีการฟื้นฟูสภาพในปีที่ใกล้สิ้นสุดระยะวิเคราะห์
        
        วิธีคำนวณ (Straight-Line Depreciation):
        
        1. หามูลค่าที่ลดลงต่อปี = (ต้นทุนฟื้นฟู × (1 - ร้อยละมูลค่าซาก/100)) / อายุใช้งาน
        2. มูลค่าซาก = ต้นทุนฟื้นฟู - (มูลค่าที่ลดลงต่อปี × จำนวนปีที่ใช้งาน)
        
        ## 5. คำแนะนำในการใช้งาน
        
        ### 📤 การ Upload ข้อมูล
        1. ดาวน์โหลด Template Excel จาก Sidebar
        2. กรอกข้อมูลต้นทุนก่อสร้างในคอลัมน์ "ต้นทุนก่อสร้าง (บาท/ตร.ม.)"
        3. อัปโหลดไฟล์กลับมา ระบบจะอัพเดทค่าอัตโนมัติ
        
        ### 💰 การใช้งาน Toggle มูลค่าซาก
        - **เปิด**: คำนวณมูลค่าซากและหักออกจากต้นทุนรวม (แนะนำ)
        - **ปิด**: ไม่คำนวณมูลค่าซาก (กรณีต้องการเปรียบเทียบต้นทุนเต็มจำนวน)
        
        ## 6. เอกสารอ้างอิง
        
        - FHWA-SA-98-079: Life-Cycle Cost Analysis in Pavement Design
        - AASHTO Guide for Design of Pavement Structures
        - NCHRP Report 703: Guide for Pavement-Type Selection
        - มาตรฐานกรมทางหลวง
        """)
    
    # ==========================================================================
    # Footer
    # ==========================================================================
    st.divider()
    st.markdown("""
    ---
    **โปรแกรมวิเคราะห์ LCCA ผิวทาง v2.1** | พัฒนาสำหรับการเรียนการสอนด้านวิศวกรรมทาง  
    ภาควิชาครุศาสตร์โยธา มหาวิทยาลัยเทคโนโลยีพระจอมเกล้าพระนครเหนือ
    
    **ฟีเจอร์ใหม่ v2.1:**
    - 📤 Upload Excel เพื่อกรอกข้อมูลต้นทุนก่อสร้าง
    - 📥 ดาวน์โหลด Template Excel
    - 🔄 เปิด/ปิดการคำนวณมูลค่าซาก
    - ⚡ Fast upload method (ไม่ต้องรอ reload)
    """)


if __name__ == "__main__":
    main()
